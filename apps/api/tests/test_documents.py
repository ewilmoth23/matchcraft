from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import fitz
import pytest
from docx import Document

from app.core.config import Settings
from app.core.errors import MatchCraftError
from app.services.documents import (
    MAX_EXTRACTED_TEXT_CHARS,
    MAX_PDF_PAGES,
    delete_stored_document,
    extract_document,
    finalize_staged_deletions,
    restore_staged_deletions,
    safe_display_filename,
    stage_stored_document_deletion,
    store_document,
    validate_document,
)


def make_pdf(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    data = document.tobytes()
    document.close()
    return data


def make_docx(text: str) -> bytes:
    document = Document()
    document.add_heading("Experience", level=1)
    document.add_paragraph(text, style="List Bullet")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def mark_docx_entry_encrypted(data: bytes, target: bytes = b"word/document.xml") -> bytes:
    """Set the ZIP encryption bit in a central entry without exposing a password fixture."""
    result = bytearray(data)
    signature = b"PK\x01\x02"
    cursor = 0
    while (header := result.find(signature, cursor)) != -1:
        name_length = int.from_bytes(result[header + 28 : header + 30], "little")
        name_start = header + 46
        if bytes(result[name_start : name_start + name_length]) == target:
            flags = int.from_bytes(result[header + 8 : header + 10], "little")
            result[header + 8 : header + 10] = (flags | 0x1).to_bytes(2, "little")
            return bytes(result)
        cursor = name_start + name_length
    raise AssertionError("DOCX fixture did not contain the target ZIP entry")


def test_pdf_validation_and_extraction(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    data = make_pdf("Built Python services for local data processing.")
    extension = validate_document(data, "resume.pdf", settings)
    extracted = extract_document(data, extension)
    assert extension == ".pdf"
    assert "Python services" in extracted.text


def test_image_only_pdf_is_reported(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    document = fitz.open()
    document.new_page()
    data = document.tobytes()
    document.close()
    extracted = extract_document(data, validate_document(data, "scan.pdf", settings))
    assert extracted.text == ""
    assert "image-only" in extracted.warnings[0]


def test_encrypted_pdf_is_rejected_explicitly(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Private résumé")
    data = document.tobytes(
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner-password",
        user_pw="user-password",
    )
    document.close()

    with pytest.raises(MatchCraftError) as error:
        extract_document(data, validate_document(data, "encrypted.pdf", settings))

    assert error.value.code == "encrypted_pdf"


def test_docx_validation_and_extraction(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    data = make_docx("Reduced turnaround time by 20%.")
    extension = validate_document(data, "resume.docx", settings)
    assert "Reduced turnaround" in extract_document(data, extension).text


@pytest.mark.parametrize(
    ("data", "filename", "code"),
    [
        (b"not pdf", "resume.pdf", "invalid_pdf"),
        (b"PK not a zip", "resume.docx", "invalid_docx"),
        (b"text", "resume.txt", "unsupported_file"),
        (b"", "resume.pdf", "empty_file"),
    ],
)
def test_invalid_documents_are_rejected(
    tmp_path: Path, data: bytes, filename: str, code: str
) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    with pytest.raises(MatchCraftError) as error:
        validate_document(data, filename, settings)
    assert error.value.code == code


def test_safe_storage_ignores_original_path(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    data = make_pdf("safe")
    stored_name = store_document(data, ".pdf", settings)
    assert safe_display_filename("../../private/resume.pdf") == "resume.pdf"
    assert "/" not in stored_name
    assert (settings.uploads_dir / stored_name).read_bytes() == data
    assert (settings.uploads_dir / stored_name).parent.resolve() == settings.uploads_dir.resolve()
    delete_stored_document(stored_name, settings)
    assert not (settings.uploads_dir / stored_name).exists()


def test_delete_rejects_path_traversal(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    settings.ensure_directories()
    outside = tmp_path / "outside.pdf"
    outside.write_bytes(b"do not delete")

    with pytest.raises(MatchCraftError) as error:
        delete_stored_document("../outside.pdf", settings)

    assert error.value.code == "unsafe_path"
    assert outside.read_bytes() == b"do not delete"


def test_staged_deletion_can_be_restored_or_finalized(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    stored_name = store_document(make_pdf("safe"), ".pdf", settings)
    original = settings.uploads_dir / stored_name

    staged = stage_stored_document_deletion(stored_name, settings)

    assert staged is not None
    assert not original.exists()
    assert staged.staged.exists()
    restore_staged_deletions([staged])
    assert original.exists()
    assert not staged.staged.exists()

    staged_again = stage_stored_document_deletion(stored_name, settings)
    assert staged_again is not None
    assert finalize_staged_deletions([staged_again]) == 0
    assert not original.exists()
    assert not staged_again.staged.exists()


def test_display_filename_removes_windows_and_posix_paths() -> None:
    assert safe_display_filename(r"C:\private\resume.pdf") == "resume.pdf"
    assert safe_display_filename("../../private/resume.pdf") == "resume.pdf"


def test_docx_archive_bomb_is_rejected_before_extraction(tmp_path: Path) -> None:
    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("word/document.xml", "0" * (2 * 1024 * 1024))
    settings = Settings(data_dir=tmp_path, provider="disabled")

    with pytest.raises(MatchCraftError) as error:
        validate_document(buffer.getvalue(), "resume.docx", settings)

    assert error.value.code == "unsafe_docx"


def test_encrypted_docx_archive_entry_is_rejected(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    data = mark_docx_entry_encrypted(make_docx("Private résumé"))

    with pytest.raises(MatchCraftError) as error:
        validate_document(data, "encrypted.docx", settings)

    assert error.value.code == "unsafe_docx"


def test_pdf_page_limit_is_enforced(tmp_path: Path) -> None:
    document = fitz.open()
    for _ in range(MAX_PDF_PAGES + 1):
        document.new_page()
    data = document.tobytes()
    document.close()
    settings = Settings(data_dir=tmp_path, provider="disabled")

    with pytest.raises(MatchCraftError) as error:
        extract_document(data, validate_document(data, "oversized.pdf", settings))

    assert error.value.code == "unsafe_pdf"


def test_docx_xml_entities_and_duplicate_members_are_rejected(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path, provider="disabled")
    entity_buffer = BytesIO()
    with ZipFile(entity_buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr(
            "word/document.xml",
            '<!DOCTYPE x [<!ENTITY secret "value">]><document>&secret;</document>',
        )
    with pytest.raises(MatchCraftError) as entity_error:
        validate_document(entity_buffer.getvalue(), "resume.docx", settings)
    assert entity_error.value.code == "unsafe_docx"

    duplicate_buffer = BytesIO()
    with (
        pytest.warns(UserWarning, match="Duplicate name"),
        ZipFile(duplicate_buffer, "w", ZIP_DEFLATED) as archive,
    ):
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("word/document.xml", "<document />")
        archive.writestr("word/document.xml", "<document />")
    with pytest.raises(MatchCraftError) as duplicate_error:
        validate_document(duplicate_buffer.getvalue(), "resume.docx", settings)
    assert duplicate_error.value.code == "unsafe_docx"


def test_extracted_text_limit_is_enforced() -> None:
    document = Document()
    text = " ".join(f"token{index}" for index in range(15_000))
    assert len(text) > MAX_EXTRACTED_TEXT_CHARS
    document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)

    with pytest.raises(MatchCraftError) as error:
        extract_document(buffer.getvalue(), ".docx")

    assert error.value.code == "document_text_too_large"
