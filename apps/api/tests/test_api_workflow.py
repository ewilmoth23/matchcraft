import json
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pytest import MonkeyPatch
from sqlalchemy import delete
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models import AnalysisScore
from app.providers.base import ProviderResponse
from app.schemas.provider import ModelAnalysisOutput


def create_completed_analysis(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> dict:
    resume = client.post("/api/v1/resumes/text", json={"text": sample_resume_text})
    assert resume.status_code == 201
    resume_id = resume.json()["id"]
    assert client.post(f"/api/v1/resumes/{resume_id}/confirm").status_code == 200
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text})
    assert job.status_code == 201
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume_id, "job_description_id": job.json()["id"]},
    )
    assert analysis.status_code == 201
    completed = client.post(
        f"/api/v1/analyses/{analysis.json()['id']}/run", json={"use_model": False}
    )
    assert completed.status_code == 200
    return completed.json()


def test_full_deterministic_api_workflow(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)
    assert result["state"] == "completed"
    assert result["model_status"] == "skipped"
    assert result["deterministic_complete"] is True
    assert len(result["scores"]) == 8
    assert result["result"]["bullet_analysis"]
    assert result["result"]["bullet_analysis"][0]["original_bullet"] in sample_resume_text
    assert any(item["requirement"] == "Python" for item in result["evidence"])
    assert (
        client.get("/api/v1/analyses").json()[0]["target_job_title"] == "Senior Software Engineer"
    )

    markdown = client.get(f"/api/v1/analyses/{result['id']}/export/markdown")
    assert markdown.status_code == 200
    assert "decision-support aid" in markdown.text
    assert "## Requirement evidence" in markdown.text
    assert "**Supporting evidence:**" in markdown.text
    assert "**Role relevance:**" in markdown.text
    assert "Confidence / source / status" in markdown.text
    assert "Talking point:" in markdown.text
    exported_json = client.get(f"/api/v1/analyses/{result['id']}/export/json")
    assert exported_json.status_code == 200
    assert "openai_api_key" not in exported_json.text
    exported = json.loads(exported_json.text)
    assert exported["overall_score"] == result["overall_score"]
    assert exported["scores"]
    assert exported["evidence"]


def test_valid_model_assisted_api_workflow(
    client: TestClient,
    sample_resume_text: str,
    sample_job_text: str,
    monkeypatch: MonkeyPatch,
) -> None:
    class FakeProvider:
        name = "ollama"
        model = "fixture-model"

        async def available(self) -> bool:
            return True

        async def analyze(self, _resume_text: str, _job_text: str) -> ProviderResponse:
            return ProviderResponse(
                ModelAnalysisOutput(
                    executive_summary="The supplied evidence supports part of the target role.",
                    responsibility_alignment=0.5,
                    transferable_experience=[
                        "Built Python and FastAPI services that processed 2 million records each week."
                    ],
                    recommendations=[],
                    interview_questions=[],
                    limitations=["Human review is required."],
                ),
                {"prompt_tokens": 100, "completion_tokens": 40},
            )

    monkeypatch.setattr(
        "app.services.model_analysis.get_provider", lambda _settings: FakeProvider()
    )
    settings_response = client.put("/api/v1/settings", json={"provider": "ollama"})
    assert settings_response.status_code == 200

    resume = client.post("/api/v1/resumes/text", json={"text": sample_resume_text}).json()
    assert client.post(f"/api/v1/resumes/{resume['id']}/confirm").status_code == 200
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text}).json()
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()
    result = client.post(f"/api/v1/analyses/{analysis['id']}/run", json={"use_model": True})

    assert result.status_code == 200
    assert result.json()["state"] == "completed"
    assert result.json()["model_status"] == "completed"
    assert result.json()["result"]["model_generated"] is True
    assert result.json()["result"]["model_responsibility_alignment"] == 0.5
    assert result.json()["result"]["model_executive_summary"].startswith("The supplied evidence")


def test_unexpected_provider_failure_preserves_deterministic_analysis(
    client: TestClient,
    sample_resume_text: str,
    sample_job_text: str,
    monkeypatch: MonkeyPatch,
) -> None:
    class BrokenProvider:
        name = "ollama"
        model = "broken-model"

        async def available(self) -> bool:
            return True

        async def analyze(self, _resume_text: str, _job_text: str) -> ProviderResponse:
            raise TypeError("malformed provider response")

    monkeypatch.setattr(
        "app.services.model_analysis.get_provider", lambda _settings: BrokenProvider()
    )
    assert client.put("/api/v1/settings", json={"provider": "ollama"}).status_code == 200

    resume = client.post("/api/v1/resumes/text", json={"text": sample_resume_text}).json()
    assert client.post(f"/api/v1/resumes/{resume['id']}/confirm").status_code == 200
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text}).json()
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()

    response = client.post(f"/api/v1/analyses/{analysis['id']}/run", json={"use_model": True})

    assert response.status_code == 200
    assert response.json()["state"] == "completed"
    assert response.json()["model_status"] == "unavailable"
    assert response.json()["deterministic_complete"] is True
    assert len(response.json()["scores"]) == 8


def test_service_revalidates_output_from_provider_adapter(
    client: TestClient,
    sample_resume_text: str,
    sample_job_text: str,
    monkeypatch: MonkeyPatch,
) -> None:
    class UnsafeProvider:
        name = "ollama"
        model = "unsafe-adapter"

        async def available(self) -> bool:
            return True

        async def analyze(self, _resume_text: str, _job_text: str) -> ProviderResponse:
            return ProviderResponse(
                ModelAnalysisOutput(
                    executive_summary=("The candidate earned a Doctorate in February 2035."),
                    responsibility_alignment=0.9,
                    transferable_experience=[],
                    recommendations=[],
                    interview_questions=[],
                    limitations=[],
                ),
                {},
            )

    monkeypatch.setattr(
        "app.services.model_analysis.get_provider", lambda _settings: UnsafeProvider()
    )
    assert client.put("/api/v1/settings", json={"provider": "ollama"}).status_code == 200
    resume = client.post("/api/v1/resumes/text", json={"text": sample_resume_text}).json()
    client.post(f"/api/v1/resumes/{resume['id']}/confirm")
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text}).json()
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()

    response = client.post(f"/api/v1/analyses/{analysis['id']}/run", json={"use_model": True})

    assert response.status_code == 200
    assert response.json()["state"] == "completed"
    assert response.json()["model_status"] == "invalid_output"
    assert response.json()["deterministic_complete"] is True
    assert "model_executive_summary" not in response.json()["result"]


def test_review_required_and_validation_errors(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    resume = client.post("/api/v1/resumes/text", json={"text": sample_resume_text}).json()
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text}).json()
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()
    response = client.post(f"/api/v1/analyses/{analysis['id']}/run", json={"use_model": False})
    assert response.status_code == 409
    assert response.json()["error"]["code"] == "resume_not_confirmed"
    blank = client.post("/api/v1/resumes/text", json={"text": "   "})
    assert blank.status_code == 422
    assert blank.json()["error"]["code"] == "validation_error"
    assert all("input" not in item for item in blank.json()["error"]["details"])


def test_history_rename_rejects_blank_names(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)
    response = client.patch(f"/api/v1/analyses/{result['id']}", json={"name": "   "})
    assert response.status_code == 422
    assert client.get(f"/api/v1/analyses/{result['id']}").json()["name"] == result["name"]


def test_settings_reject_secret_bearing_provider_url(client: TestClient) -> None:
    response = client.put(
        "/api/v1/settings",
        json={"provider": "openai_compatible", "provider_url": "https://user:secret@api.test/v1"},
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "validation_error"
    assert "secret" not in response.text


def test_settings_persist_separate_local_and_remote_defaults(client: TestClient) -> None:
    response = client.put(
        "/api/v1/settings",
        json={
            "provider": "local_first",
            "local_model": "qwen3.5:9b",
            "local_provider_url": "http://localhost:11434",
            "remote_model": "gpt-5.6-sol",
            "remote_provider_url": "https://api.openai.com/v1",
            "openai_reasoning_effort": "low",
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["provider"] == "local_first"
    assert body["local_model"] == "qwen3.5:9b"
    assert body["remote_model"] == "gpt-5.6-sol"
    assert body["local_provider_url"] == "http://localhost:11434"
    assert body["remote_provider_url"] == "https://api.openai.com/v1"
    assert body["remote_api_key_configured"] is False
    assert body["remote_fallback_configured"] is False
    assert "api_key" not in body


def test_resume_edit_invalidates_stale_analysis_outputs(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)
    replacement = (
        "Jordan Rivera\nEXPERIENCE\n"
        "• Maintained internal documentation and reviewed support procedures.\n"
        "SKILLS\nDocumentation"
    )

    updated = client.put(
        f"/api/v1/resumes/{result['resume_id']}", json={"extracted_text": replacement}
    )

    assert updated.status_code == 200
    assert updated.json()["confirmed"] is False
    invalidated = client.get(f"/api/v1/analyses/{result['id']}").json()
    assert invalidated["state"] == "draft"
    assert invalidated["deterministic_complete"] is False
    assert invalidated["overall_score"] is None
    assert invalidated["model_status"] == "not_requested"
    assert invalidated["result"] == {}
    assert invalidated["scores"] == []
    assert invalidated["evidence"] == []
    assert invalidated["recommendations"] == []
    assert invalidated["interview_questions"] == []
    assert client.get(f"/api/v1/analyses/{result['id']}/export/json").status_code == 409

    assert client.post(f"/api/v1/resumes/{result['resume_id']}/confirm").status_code == 200
    assert client.get(f"/api/v1/analyses/{result['id']}").json()["state"] == "ready"


def test_job_edit_invalidates_stale_analysis_outputs(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)
    replacement = (
        "Documentation Specialist\nRequired Qualifications\n"
        "• Technical writing is required.\nResponsibilities\n"
        "• Maintain internal procedures and support documentation."
    )

    updated = client.put(
        f"/api/v1/job-descriptions/{result['job_description_id']}", json={"text": replacement}
    )

    assert updated.status_code == 200
    invalidated = client.get(f"/api/v1/analyses/{result['id']}").json()
    assert invalidated["state"] == "ready"
    assert invalidated["deterministic_complete"] is False
    assert invalidated["scores"] == []
    assert invalidated["evidence"] == []
    assert client.get(f"/api/v1/analyses/{result['id']}/export/markdown").status_code == 409


def test_failed_rerun_rolls_back_partial_result_replacement(
    client: TestClient,
    sample_resume_text: str,
    sample_job_text: str,
    monkeypatch: MonkeyPatch,
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)
    original_scores = result["scores"]

    def fail_after_deleting_scores(db, analysis) -> None:  # type: ignore[no-untyped-def]
        db.execute(delete(AnalysisScore).where(AnalysisScore.analysis_id == analysis.id))
        analysis.overall_score = 1
        raise RuntimeError("forced deterministic failure")

    monkeypatch.setattr(
        "app.api.v1.analyses.execute_deterministic_analysis", fail_after_deleting_scores
    )

    response = client.post(f"/api/v1/analyses/{result['id']}/rerun", json={"use_model": False})

    assert response.status_code == 500
    preserved = client.get(f"/api/v1/analyses/{result['id']}").json()
    assert preserved["state"] == "failed"
    assert preserved["deterministic_complete"] is True
    assert preserved["overall_score"] == result["overall_score"]
    assert preserved["scores"] == original_scores


def test_bullet_rewrite_uses_placeholder_and_rejects_unknown_source(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)
    bullet = "Created data-quality dashboards and documented remediation procedures."
    rewrite = client.post(
        f"/api/v1/analyses/{result['id']}/bullet-rewrite",
        json={"original_bullet": bullet},
    )
    assert rewrite.status_code == 200
    assert "[insert verified outcome]" in rewrite.json()["suggested_bullet"]
    assert rewrite.json()["confirmation_required"] is True
    assert rewrite.json()["model_generated"] is False
    rejected = client.post(
        f"/api/v1/analyses/{result['id']}/bullet-rewrite",
        json={"original_bullet": "Invented accomplishment"},
    )
    assert rejected.status_code == 422


def test_bullet_rewrite_preserves_existing_quantified_outcome(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)
    bullet = (
        "Reduced data-validation failures by 32% by adding typed contracts and automated tests."
    )
    assert bullet in sample_resume_text

    rewrite = client.post(
        f"/api/v1/analyses/{result['id']}/bullet-rewrite",
        json={"original_bullet": bullet},
    )

    assert rewrite.status_code == 200
    payload = rewrite.json()
    assert payload["suggested_bullet"] == bullet
    assert "[insert verified" not in payload["suggested_bullet"]
    assert payload["confirmation_required"] is True
    assert payload["model_generated"] is False


def test_deletion_removes_analysis_sources(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)
    response = client.delete(f"/api/v1/analyses/{result['id']}")
    assert response.status_code == 204
    assert client.get(f"/api/v1/analyses/{result['id']}").status_code == 404
    assert client.get(f"/api/v1/resumes/{result['resume_id']}").status_code == 404
    assert client.get(f"/api/v1/job-descriptions/{result['job_description_id']}").status_code == 404


def test_direct_resume_deletion_removes_orphaned_related_records(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)

    assert client.delete(f"/api/v1/resumes/{result['resume_id']}").status_code == 204

    assert client.get(f"/api/v1/analyses/{result['id']}").status_code == 404
    assert client.get(f"/api/v1/job-descriptions/{result['job_description_id']}").status_code == 404


def test_direct_job_deletion_removes_orphaned_related_records(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    result = create_completed_analysis(client, sample_resume_text, sample_job_text)

    assert (
        client.delete(f"/api/v1/job-descriptions/{result['job_description_id']}").status_code == 204
    )

    assert client.get(f"/api/v1/analyses/{result['id']}").status_code == 404
    assert client.get(f"/api/v1/resumes/{result['resume_id']}").status_code == 404


def test_deletion_removes_uploaded_file_and_related_records(
    client: TestClient, sample_job_text: str
) -> None:
    document = Document()
    document.add_paragraph(
        "Jordan Rivera built reliable Python services for public-sector data processing."
    )
    buffer = BytesIO()
    document.save(buffer)
    resume_response = client.post(
        "/api/v1/resumes/upload",
        files={
            "file": (
                "resume.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert resume_response.status_code == 201
    resume = resume_response.json()
    assert client.post(f"/api/v1/resumes/{resume['id']}/confirm").status_code == 200
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text}).json()
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()
    test_settings = client.app.dependency_overrides[get_settings]()
    data_dir = Path(test_settings.data_dir)
    uploaded_files = list((data_dir / "uploads").iterdir())
    assert len(uploaded_files) == 1

    assert client.delete(f"/api/v1/analyses/{analysis['id']}").status_code == 204

    assert not uploaded_files[0].exists()
    assert client.get(f"/api/v1/resumes/{resume['id']}").status_code == 404
    assert client.get(f"/api/v1/job-descriptions/{job['id']}").status_code == 404


def test_delete_commit_failure_restores_staged_upload(
    client: TestClient,
    sample_job_text: str,
    monkeypatch: MonkeyPatch,
) -> None:
    document = Document()
    document.add_paragraph(
        "Jordan Rivera built reliable Python services for public-sector data processing."
    )
    buffer = BytesIO()
    document.save(buffer)
    resume = client.post(
        "/api/v1/resumes/upload",
        files={"file": ("resume.docx", buffer.getvalue())},
    ).json()
    client.post(f"/api/v1/resumes/{resume['id']}/confirm")
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text}).json()
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()
    test_settings = client.app.dependency_overrides[get_settings]()
    uploaded_file = next(test_settings.uploads_dir.iterdir())

    def fail_commit(_session: Session) -> None:
        raise RuntimeError("forced commit failure")

    monkeypatch.setattr(Session, "commit", fail_commit)

    with pytest.raises(RuntimeError, match="forced commit failure"):
        client.delete(f"/api/v1/analyses/{analysis['id']}")

    assert uploaded_file.exists()
    assert not list(test_settings.uploads_dir.glob(".matchcraft-delete-*"))
    assert client.get(f"/api/v1/analyses/{analysis['id']}").status_code == 200


def test_health_reports_ai_unavailable(client: TestClient) -> None:
    health = client.get("/api/v1/health")
    assert health.status_code == 200
    assert health.json()["database"] == "healthy"
    assert health.json()["ai_features"] == "unavailable"


@pytest.mark.asyncio
async def test_health_degrades_when_database_is_unavailable() -> None:
    from app.api.v1.health import health

    class BrokenSession:
        def execute(self, _statement: object) -> None:
            raise RuntimeError("database unavailable")

    result = await health(BrokenSession())  # type: ignore[arg-type]
    assert result.status == "degraded"
    assert result.database == "unavailable"
    assert result.deterministic_analysis == "unavailable"


def test_routine_logs_do_not_include_document_bodies(
    client: TestClient,
    caplog: pytest.LogCaptureFixture,
    capsys: pytest.CaptureFixture[str],
) -> None:
    resume_marker = "PRIVATE_RESUME_MARKER_8f42"
    job_marker = "PRIVATE_JOB_MARKER_91ad"
    resume_text = f"Jordan Rivera\nEXPERIENCE\n• Built Python services.\n{resume_marker}"
    job_text = f"Engineer\nRequired Qualifications\n• Python required.\n{job_marker}"
    resume = client.post("/api/v1/resumes/text", json={"text": resume_text}).json()
    client.post(f"/api/v1/resumes/{resume['id']}/confirm")
    job = client.post("/api/v1/job-descriptions", json={"text": job_text}).json()
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()
    client.post(f"/api/v1/analyses/{analysis['id']}/run", json={"use_model": False})
    path_marker = "PRIVATE_PATH_MARKER_c773"
    request_id_marker = "PRIVATE_REQUEST_ID_MARKER_a816"
    client.get(
        f"/api/v1/analyses/{path_marker}",
        headers={"X-Request-ID": request_id_marker},
    )

    application_logs = capsys.readouterr().out
    captured_logs = caplog.text + application_logs
    assert resume_marker not in captured_logs
    assert job_marker not in captured_logs
    assert path_marker not in application_logs
    assert request_id_marker not in application_logs
