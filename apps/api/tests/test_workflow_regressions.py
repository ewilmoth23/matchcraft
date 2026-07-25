"""Workflow-level regression tests from the 2026-07-24 expert review."""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pytest import MonkeyPatch
from sqlalchemy.orm import Session

from app.core.config import get_settings
from tests.test_api_workflow import create_completed_analysis


def test_bullet_rewrite_requires_a_confirmed_resume(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    """The rewrite endpoint could send unreviewed résumé text to a provider."""
    resume = client.post("/api/v1/resumes/text", json={"text": sample_resume_text}).json()
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text}).json()
    client.post(f"/api/v1/resumes/{resume['id']}/confirm")
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()

    bullet = next(line for line in sample_resume_text.splitlines() if line.strip().startswith("•"))
    bullet = bullet.strip().lstrip("• ")

    # Editing the résumé clears confirmation and must also block a rewrite.
    edited = client.put(
        f"/api/v1/resumes/{resume['id']}",
        json={"extracted_text": sample_resume_text + "\nAdditional reviewed detail."},
    )
    assert edited.status_code == 200
    assert edited.json()["confirmed"] is False
    blocked = client.post(
        f"/api/v1/analyses/{analysis['id']}/bullet-rewrite",
        json={"original_bullet": bullet},
    )
    assert blocked.status_code == 409
    assert blocked.json()["error"]["code"] == "resume_not_confirmed"


def test_markdown_export_includes_transferable_experience(
    client: TestClient, sample_resume_text: str, sample_job_text: str
) -> None:
    """The JSON export carried this section and the UI renders it; Markdown omitted it."""
    # A purpose-built pairing rather than the shared samples: recalibrating the matcher
    # moved those to fully supported, which silently stopped exercising this export path.
    # The case is constructed so the requirement is genuinely adjacent, not equivalent.
    resume = (
        "Ada Fletcher\nada.fletcher@example.test\n\n"
        "EXPERIENCE\nOperations Analyst | Meridian Freight | 2020 - Present\n"
        "• Coordinated regional distribution schedules for eleven partner depots.\n\n"
        "SKILLS\nExcel, Jira\n\nEDUCATION\nBA Logistics | Fairview College | 2019\n"
    )
    job = (
        "Supply Chain Planner\nRequirements\n"
        "• Manage national distribution planning for retail partner networks nationwide.\n"
    )
    analysis = create_completed_analysis(client, resume, job)
    transferable = analysis["result"].get("transferable_experience")
    assert transferable, "the constructed pairing no longer produces a transferable finding"
    markdown = client.get(f"/api/v1/analyses/{analysis['id']}/export/markdown").text
    assert "Potentially transferable experience" in markdown
    for item in transferable:
        assert item in markdown


@pytest.mark.parametrize("resource", ["job-descriptions", "resumes", "analyses"])
def test_delete_commit_failure_restores_staged_upload(
    client: TestClient,
    sample_job_text: str,
    monkeypatch: MonkeyPatch,
    resource: str,
) -> None:
    """All three entry points share one transaction; each must restore staged files.

    Only the analysis variant was covered before, and the job variant staged files in a
    different order.
    """
    document = Document()
    document.add_paragraph("Dana Chen built reliable Python services for public data processing.")
    buffer = BytesIO()
    document.save(buffer)
    resume = client.post(
        "/api/v1/resumes/upload", files={"file": ("resume.docx", buffer.getvalue())}
    ).json()
    client.post(f"/api/v1/resumes/{resume['id']}/confirm")
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text}).json()
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()
    target = {"job-descriptions": job["id"], "resumes": resume["id"], "analyses": analysis["id"]}[
        resource
    ]
    test_settings = client.app.dependency_overrides[get_settings]()  # type: ignore[attr-defined]
    uploaded_file = next(Path(test_settings.uploads_dir).iterdir())

    def fail_commit(_session: Session) -> None:
        raise RuntimeError("forced commit failure")

    monkeypatch.setattr(Session, "commit", fail_commit)
    with pytest.raises(RuntimeError, match="forced commit failure"):
        client.delete(f"/api/v1/{resource}/{target}")

    assert uploaded_file.exists()
    assert not list(Path(test_settings.uploads_dir).glob(".matchcraft-delete-*"))


@pytest.mark.parametrize("resource", ["job-descriptions", "resumes", "analyses"])
def test_failure_during_staging_restores_the_already_staged_upload(
    client: TestClient,
    sample_job_text: str,
    monkeypatch: MonkeyPatch,
    resource: str,
) -> None:
    """A raise *between* two staged files must not strand the first one.

    The upload is renamed before the exports are staged. If the staged list is built
    inside the staging step rather than owned by the transaction, a failure there leaves
    a renamed upload that nothing can restore — and the startup sweep then deletes it,
    while the résumé row survives pointing at a file that no longer exists.
    """
    document = Document()
    document.add_paragraph("Dana Chen built reliable Python services for public data processing.")
    buffer = BytesIO()
    document.save(buffer)
    resume = client.post(
        "/api/v1/resumes/upload", files={"file": ("resume.docx", buffer.getvalue())}
    ).json()
    client.post(f"/api/v1/resumes/{resume['id']}/confirm")
    job = client.post("/api/v1/job-descriptions", json={"text": sample_job_text}).json()
    analysis = client.post(
        "/api/v1/analyses",
        json={"resume_id": resume["id"], "job_description_id": job["id"]},
    ).json()
    target = {"job-descriptions": job["id"], "resumes": resume["id"], "analyses": analysis["id"]}[
        resource
    ]
    test_settings = client.app.dependency_overrides[get_settings]()  # type: ignore[attr-defined]
    uploaded_file = next(Path(test_settings.uploads_dir).iterdir())

    def fail_after_the_upload_is_staged(*_args: object, **_kwargs: object) -> None:
        raise OSError("export directory is unreadable")

    monkeypatch.setattr(
        "app.services.deletion.stage_export_deletions", fail_after_the_upload_is_staged
    )
    with pytest.raises(OSError, match="export directory is unreadable"):
        client.delete(f"/api/v1/{resource}/{target}")

    assert uploaded_file.exists(), "the staged upload was not restored"
    assert not list(Path(test_settings.uploads_dir).glob(".matchcraft-delete-*"))
    assert client.get(f"/api/v1/resumes/{resume['id']}").status_code == 200


def test_failure_between_two_export_files_restores_both(
    client: TestClient,
    sample_resume_text: str,
    sample_job_text: str,
    monkeypatch: MonkeyPatch,
) -> None:
    """`stage_export_deletions` must append to the caller's list, not its own.

    A raise between two export files would otherwise strand the first one: the caller's
    restore path never sees it, and the startup sweep removes it.
    """
    analysis = create_completed_analysis(client, sample_resume_text, sample_job_text)
    test_settings = client.app.dependency_overrides[get_settings]()  # type: ignore[attr-defined]
    exports = Path(test_settings.exports_dir)
    first = exports / f"{analysis['id']}-report.md"
    second = exports / f"{analysis['id']}-report.json"
    first.write_text("# saved export", encoding="utf-8")
    second.write_text("{}", encoding="utf-8")

    from app.services import documents

    original = documents._stage_file_deletion
    calls = {"count": 0}

    def fail_on_the_second_file(*args: object, **kwargs: object) -> object:
        calls["count"] += 1
        if calls["count"] == 2:
            raise OSError("the second export could not be staged")
        return original(*args, **kwargs)  # type: ignore[arg-type]

    monkeypatch.setattr(documents, "_stage_file_deletion", fail_on_the_second_file)
    with pytest.raises(OSError, match="the second export could not be staged"):
        client.delete(f"/api/v1/analyses/{analysis['id']}")

    assert calls["count"] >= 2, "the probe never reached the second file"
    assert first.exists() and second.exists(), "a staged export was not restored"
    assert not list(exports.glob(".matchcraft-delete-*"))


def test_startup_sweeps_staged_deletion_residue(tmp_path: Path) -> None:
    """A crash between staging and finalizing used to strand the upload forever.

    Drives the real lifespan hook rather than the helper, so the wiring is covered too.
    """
    from app.core.config import Settings
    from app.main import app
    from app.main import settings as application_settings

    isolated = Settings(
        env="test",
        database_url=f"sqlite:///{tmp_path / 'lifespan.db'}",
        data_dir=tmp_path / "data",
        provider="disabled",
    )
    isolated.ensure_directories()
    residue = isolated.uploads_dir / ".matchcraft-delete-abc123"
    residue.write_bytes(b"orphan")
    kept = isolated.uploads_dir / "keep.pdf"
    kept.write_bytes(b"%PDF-1.4")

    original_dir = application_settings.data_dir
    try:
        application_settings.data_dir = isolated.data_dir
        with TestClient(app):
            pass
    finally:
        application_settings.data_dir = original_dir

    assert not residue.exists()
    assert kept.exists()
