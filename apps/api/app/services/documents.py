import os
import re
import uuid
import zipfile
from contextlib import suppress
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import fitz
from docx import Document

from app.core.config import Settings
from app.core.errors import MatchCraftError

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MIME_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
MAX_DOCX_ARCHIVE_MEMBERS = 2_000
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_DOCX_MEMBER_BYTES = 20 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 200
MAX_EXTRACTED_TEXT_CHARS = 100_000
MAX_PDF_PAGES = 250


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    warnings: list[str]
    media_type: str


@dataclass(frozen=True)
class StagedFileDeletion:
    original: Path
    staged: Path


def safe_display_filename(filename: str | None) -> str:
    name = re.split(r"[\\/]", filename or "resume")[-1]
    name = re.sub(r"[\x00-\x1f\x7f]", "", name).strip()
    return name[:255] or "resume"


def validate_document(data: bytes, filename: str | None, settings: Settings) -> str:
    if not data:
        raise MatchCraftError("empty_file", "The uploaded résumé is empty.", 422)
    if len(data) > settings.max_upload_bytes:
        raise MatchCraftError(
            "file_too_large",
            f"The file exceeds the {settings.max_upload_bytes // (1024 * 1024)} MB upload limit.",
            413,
        )
    extension = Path(safe_display_filename(filename)).suffix.casefold()
    if extension not in ALLOWED_EXTENSIONS:
        raise MatchCraftError("unsupported_file", "Upload a PDF or DOCX résumé.", 415)
    if extension == ".pdf" and not data.startswith(b"%PDF-"):
        raise MatchCraftError(
            "invalid_pdf", "The file does not contain a valid PDF signature.", 422
        )
    if extension == ".docx":
        if not data.startswith(b"PK"):
            raise MatchCraftError(
                "invalid_docx", "The file does not contain a valid DOCX archive.", 422
            )
        try:
            with zipfile.ZipFile(BytesIO(data)) as archive:
                members = archive.infolist()
                names = {item.filename for item in members}
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise MatchCraftError(
                        "invalid_docx", "The archive is not a DOCX document.", 422
                    )
                if len(members) > MAX_DOCX_ARCHIVE_MEMBERS:
                    raise MatchCraftError(
                        "unsafe_docx", "The DOCX archive contains too many entries.", 422
                    )
                total_uncompressed = sum(item.file_size for item in members)
                if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES or any(
                    item.file_size > MAX_DOCX_MEMBER_BYTES for item in members
                ):
                    raise MatchCraftError(
                        "unsafe_docx",
                        "The expanded DOCX archive is too large to process safely.",
                        422,
                    )
                if any(item.flag_bits & 0x1 for item in members):
                    raise MatchCraftError(
                        "unsafe_docx", "Encrypted DOCX archive entries are not supported.", 422
                    )
                if any(
                    item.file_size > 0
                    and item.file_size / max(1, item.compress_size) > MAX_DOCX_COMPRESSION_RATIO
                    for item in members
                ):
                    raise MatchCraftError(
                        "unsafe_docx",
                        "The DOCX archive has an unsafe compression ratio.",
                        422,
                    )
                if len(names) != len(members):
                    raise MatchCraftError(
                        "unsafe_docx", "The DOCX archive contains duplicate entries.", 422
                    )
                # python-docx parses styles, numbering, relationships, and more — not just
                # word/document.xml — so every XML part has to be screened.
                for member in members:
                    if not member.filename.casefold().endswith((".xml", ".rels")):
                        continue
                    upper_xml = archive.read(member.filename).upper()
                    if b"<!DOCTYPE" in upper_xml or b"<!ENTITY" in upper_xml:
                        raise MatchCraftError(
                            "unsafe_docx",
                            "DOCX files containing document type or entity declarations are not supported.",
                            422,
                        )
        except zipfile.BadZipFile as exc:
            raise MatchCraftError("invalid_docx", "The DOCX archive is corrupt.", 422) from exc
    return extension


def extract_document(data: bytes, extension: str) -> ExtractedDocument:
    if extension == ".pdf":
        return _extract_pdf(data)
    if extension == ".docx":
        return _extract_docx(data)
    raise MatchCraftError("unsupported_file", "Upload a PDF or DOCX résumé.", 415)


def _extract_pdf(data: bytes) -> ExtractedDocument:
    try:
        with fitz.open(stream=data, filetype="pdf") as document:
            if document.needs_pass:
                raise MatchCraftError(
                    "encrypted_pdf",
                    "Password-protected PDF files are not supported; unlock the file or paste its text.",
                    422,
                )
            if document.page_count > MAX_PDF_PAGES:
                raise MatchCraftError(
                    "unsafe_pdf",
                    f"PDF files may contain at most {MAX_PDF_PAGES} pages.",
                    422,
                )
            pages = [page.get_text("text", sort=True).strip() for page in document]
    except MatchCraftError:
        raise
    except Exception as exc:
        raise MatchCraftError(
            "corrupt_pdf", "The PDF could not be read and may be corrupt.", 422
        ) from exc
    text = "\n\n".join(page for page in pages if page).strip()
    _validate_extracted_text_size(text)
    warnings: list[str] = []
    if not text:
        warnings.append(
            "No selectable text was found. This may be an image-only PDF; use OCR elsewhere or paste the résumé text."
        )
    elif len(text) < 80:
        warnings.append(
            "Very little text was extracted. Review the text carefully before confirming it."
        )
    return ExtractedDocument(text, warnings, MIME_TYPES[".pdf"])


def _extract_docx(data: bytes) -> ExtractedDocument:
    try:
        document = Document(BytesIO(data))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                prefix = "• " if paragraph.style and "List" in paragraph.style.name else ""
                blocks.append(prefix + text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
    except Exception as exc:
        raise MatchCraftError(
            "corrupt_docx", "The DOCX could not be read and may be corrupt.", 422
        ) from exc
    text = "\n".join(blocks).strip()
    _validate_extracted_text_size(text)
    warnings = [] if text else ["No readable text was found in the DOCX document."]
    return ExtractedDocument(text, warnings, MIME_TYPES[".docx"])


def _validate_extracted_text_size(text: str) -> None:
    if len(text) > MAX_EXTRACTED_TEXT_CHARS:
        raise MatchCraftError(
            "document_text_too_large",
            f"Extracted résumé text may contain at most {MAX_EXTRACTED_TEXT_CHARS:,} characters.",
            422,
        )


def store_document(data: bytes, extension: str, settings: Settings) -> str:
    settings.ensure_directories()
    stored_name = f"{uuid.uuid4().hex}{extension}"
    base = settings.uploads_dir.resolve()
    destination = (base / stored_name).resolve()
    if destination.parent != base:
        raise MatchCraftError("unsafe_path", "Could not create a safe storage path.", 500)
    # Create with restrictive permissions rather than widening then narrowing them,
    # which briefly exposed résumé uploads at the process umask.
    descriptor = os.open(destination, os.O_WRONLY | os.O_CREAT | os.O_EXCL, 0o600)
    with os.fdopen(descriptor, "wb") as handle:
        handle.write(data)
    return stored_name


def sweep_staged_deletions(settings: Settings) -> int:
    """Remove staged-deletion residue left by an interrupted delete.

    A crash between staging and finalizing otherwise strands the original upload on
    disk under a hidden name with nothing to reclaim it.
    """
    base = settings.uploads_dir
    if not base.exists():
        return 0
    removed = 0
    for candidate in base.glob(".matchcraft-delete-*"):
        with suppress(OSError):
            if candidate.is_file():
                candidate.unlink()
                removed += 1
    return removed


def delete_stored_document(stored_filename: str | None, settings: Settings) -> None:
    if not stored_filename:
        return
    base = settings.uploads_dir.resolve()
    candidate = (base / stored_filename).resolve()
    if candidate.parent != base:
        raise MatchCraftError("unsafe_path", "Refusing to delete an unsafe storage path.", 500)
    if candidate.exists() and candidate.is_file():
        candidate.unlink()


def stage_stored_document_deletion(
    stored_filename: str | None, settings: Settings
) -> StagedFileDeletion | None:
    if not stored_filename:
        return None
    base = settings.uploads_dir.resolve()
    candidate = (base / stored_filename).resolve()
    if candidate.parent != base:
        raise MatchCraftError("unsafe_path", "Refusing to delete an unsafe storage path.", 500)
    return _stage_file_deletion(candidate, base)


def stage_export_deletions(
    analysis_ids: list[str],
    settings: Settings,
    into: list[StagedFileDeletion] | None = None,
) -> list[StagedFileDeletion]:
    """Stage export files for deletion, appending to `into` as each rename happens.

    Pass the caller's list so a failure part-way through still leaves every file that
    was already renamed visible to the caller's restore path.
    """
    staged: list[StagedFileDeletion] = [] if into is None else into
    base = settings.exports_dir.resolve()
    if not base.exists():
        return staged
    for analysis_id in analysis_ids:
        if not re.fullmatch(r"[A-Fa-f0-9-]{36}", analysis_id):
            raise MatchCraftError("unsafe_path", "Refusing to delete an unsafe export path.", 500)
        for candidate in base.glob(f"{analysis_id}-*"):
            deletion = _stage_file_deletion(candidate.resolve(), base)
            if deletion:
                staged.append(deletion)
    return staged


def restore_staged_deletions(deletions: list[StagedFileDeletion]) -> None:
    for deletion in reversed(deletions):
        if deletion.staged.exists():
            deletion.staged.replace(deletion.original)


def finalize_staged_deletions(deletions: list[StagedFileDeletion]) -> int:
    failures = 0
    for deletion in deletions:
        try:
            if deletion.staged.exists():
                deletion.staged.unlink()
        except OSError:
            failures += 1
    return failures


def _stage_file_deletion(candidate: Path, base: Path) -> StagedFileDeletion | None:
    if candidate.parent != base:
        raise MatchCraftError("unsafe_path", "Refusing to delete an unsafe storage path.", 500)
    if not candidate.exists():
        return None
    if not candidate.is_file():
        raise MatchCraftError("unsafe_path", "Refusing to delete a non-file storage path.", 500)
    staged = base / f".matchcraft-delete-{uuid.uuid4().hex}"
    candidate.replace(staged)
    return StagedFileDeletion(original=candidate, staged=staged)
